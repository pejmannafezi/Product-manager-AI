"""Extract plain text from uploaded project documents (docx, pdf, txt, md)."""
import io
import re


class DocParseError(Exception):
    pass


def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".docx"):
        text = _from_docx(data)
    elif name.endswith(".pdf"):
        text = _from_pdf(data)
    elif name.endswith((".txt", ".md")):
        text = data.decode("utf-8", errors="replace")
    else:
        raise DocParseError(f"Unsupported file type: {filename}. Use .docx, .pdf, .txt, or .md")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) < 50:
        raise DocParseError(
            "Could not extract readable text from this file. "
            "If it is a scanned PDF, export it as text or docx first."
        )
    return text


def _from_docx(data: bytes) -> str:
    import docx
    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)
