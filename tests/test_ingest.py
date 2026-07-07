import docx

from scripts.ingest_docx import CHAPTER_RE, ingest, split_chapters


def test_chapter_regex():
    assert CHAPTER_RE.match("Chapter 1: The AI Product Manager Role")
    assert CHAPTER_RE.match("Chapter 23 : Comprehensive Glossary")
    assert not CHAPTER_RE.match("In Chapter 3 we discussed data")


def test_split_chapters():
    blocks = [
        ("Normal", "Some intro text"),
        ("Normal", "Chapter 1: First"),
        ("Normal", "Body one"),
        ("Normal", "Chapter 2: Second"),
        ("Heading 2", "A Section"),
        ("Normal", "Body two"),
    ]
    chapters = split_chapters(blocks)
    numbers = [c[0] for c in chapters]
    assert numbers == [0, 1, 2]
    assert chapters[2][1] == "Second"


def test_ingest_small_docx(conn, tmp_path):
    doc = docx.Document()
    doc.add_paragraph("Chapter 1: Alpha")
    doc.add_paragraph("Rollback plans matter for AI products in production systems.")
    doc.add_paragraph("Chapter 2: Beta")
    doc.add_paragraph("Human review and oversight keep clinicians in control.")
    doc.add_paragraph("Chapter 3: Gamma Glossary")
    doc.add_paragraph("Model Drift")
    doc.add_paragraph("Gradual performance change over time as data distributions shift away from training data.")
    path = tmp_path / "mini.docx"
    doc.save(str(path))

    count = ingest(path, replace=True)
    assert count == 3
    rows = conn.execute("SELECT COUNT(*) FROM kb_chapters WHERE number > 0").fetchone()[0]
    assert rows == 3
    # FTS works
    hits = conn.execute(
        "SELECT rowid FROM kb_fts WHERE kb_fts MATCH '\"rollback\"'").fetchall()
    assert hits
    # glossary harvested from the glossary chapter
    term = conn.execute("SELECT * FROM glossary WHERE term='Model Drift'").fetchone()
    assert term and term["chapter_id"] is not None
